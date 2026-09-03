#!/usr/bin/env python3
"""
docx_out.py — render the assembled term markdown into the school's print format: A4 landscape,
two columns, one document per class-term.

Kept separate from build_term_doc.py so the layout can be tuned without touching the rules, and so a
teacher's edition can reuse it later. Layout notes:

  * landscape A4 (29.7 × 21.0 cm) with 1.4 cm margins — the notes are wide, short-line material;
  * two columns, 0.7 cm apart with a rule between them; a column break is left to Word;
  * body 10.5 pt Calibri, 1.05 line spacing, 2 pt after a paragraph — this is a book the child writes
    in, so nothing is padded;
  * markdown **bold** / *italic* become real runs; "•" bullets and "A)" options get hanging indents so
    a wrapped line never starts under the bullet.

    from docx_out import render; render(markdown_text, "out.docx")
"""
import re
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).resolve().parent))
from book_layout import (asset_path, BULLET, IMAGE_WIDTH_CM, NUMBERED, OPTION, flow_of, is_week, join_wrapped, numbered_flow, plan)   # noqa: E402
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY = 10.5
GREY = RGBColor(0x55, 0x55, 0x55)
NAVY = RGBColor(0x14, 0x2C, 0x5A)
BULLET = re.compile(r"^\s*(?:[•·]\s+|-\s+)")
NUMBERED = re.compile(r"^\s*(\d{1,2})([.)])\s+(.*)$")
OPTION = re.compile(r"^\s*([A-D])\)\s+(.*)$")
INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def _runs(p, text, size=BODY, italic=False, bold=False, colour=None):
    """write markdown inline emphasis into real runs"""
    for part in INLINE.split(text):
        if not part:
            continue
        b, i, t = bold, italic, part
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            t, b = part[2:-2], True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            t, i = part[1:-1], True
        r = p.add_run(t)
        r.font.size = Pt(size)
        r.bold = b
        r.italic = i
        if colour is not None:
            r.font.color.rgb = colour
    return p


def _style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(BODY)
    for name in HEAD_STYLE.values():                 # headings carry a style only so the contents can
        try:                                          # be built from it — never let it change the look
            h = doc.styles[name]
        except KeyError:
            continue
        h.font.name = "Calibri"
        h.font.size = Pt(BODY)
        h.font.bold = True
        h.font.color.rgb = NAVY
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(3)
        h.paragraph_format.keep_with_next = True
        h.paragraph_format.line_spacing = 1.05      # the body's spacing, not the template's
    pf = st.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.05


def _landscape_two_col(doc, cols=2, gap_cm=0.7, rule=True):
    for s in doc.sections:
        s.orientation = WD_ORIENT.LANDSCAPE
        s.page_width, s.page_height = Cm(29.7), Cm(21.0)
        s.left_margin = s.right_margin = Cm(1.4)
        s.top_margin = s.bottom_margin = Cm(1.3)
        sectPr = s._sectPr
        el = sectPr.find(qn("w:cols"))
        if el is None:
            el = OxmlElement("w:cols")
            sectPr.append(el)
        el.set(qn("w:num"), str(cols))
        el.set(qn("w:space"), str(int(gap_cm * 567)))
        el.set(qn("w:sep"), "1" if rule else "0")
        el.set(qn("w:equalWidth"), "1")


def _indent(p, left, hang):
    pf = p.paragraph_format
    pf.left_indent = Cm(left)
    pf.first_line_indent = Cm(-hang)
    pf.space_after = Pt(1)


# A heading gets a real Word style so the contents page can be built from it: level 1 for a subject or
# a practice paper, level 2 for a week (or a stream), level 3 for anything the pupil does not navigate
# to (Section A/B/C of a paper, "What this book holds"). The field reads levels 1-2 only.
HEAD_STYLE = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}


def _heading(doc, text, level, size, colour=NAVY, caps=False, before=8, after=3, toc=0):
    p = doc.add_paragraph()
    if toc:                                   # only a styled heading is picked up by the TOC field
        try:
            p.style = doc.styles[HEAD_STYLE[toc]]
        except KeyError:
            pass
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.keep_with_next = True
    _runs(p, text.upper() if caps else text, size=size, bold=True, colour=colour)
    return p


def _toc_styles(doc, right_cm, size=BODY - 0.5):
    """TOC 1/TOC 2 as Word will use them: a dot leader to a right tab at the column edge."""
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    for lvl, (indent, bold) in enumerate([(0.0, True), (0.45, False)], 1):
        name = f"TOC {lvl}"
        try:
            st = doc.styles[name]
        except KeyError:
            st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = doc.styles["Normal"]
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = NAVY if bold else None
        pf = st.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(2 if lvl == 1 else 0)
        pf.line_spacing = 1.0
        pf.left_indent = Cm(indent)
        pf.tab_stops.add_tab_stop(Cm(right_cm), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


def _toc(doc, entries, right_cm, open_field=True, close_field=True):
    """The Contents page: a real Word TOC field, with our own list cached inside it so the page is
    still readable in a viewer that never updates fields. An entry that carries a stamped page number
    (tools/book_pages.py read it off a real render) prints it after a dot leader."""
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    _toc_styles(doc, right_cm)
    FIELD = ' TOC \\o "1-2" \\h \\z \\u '
    entries = [(e + (None,))[:3] if len(e) == 2 else e[:3] for e in entries]
    for i, (lvl, text, page) in enumerate(entries):
        p = doc.add_paragraph()          # deliberately not styled: Word restyles these on update
        pf = p.paragraph_format
        pf.left_indent = Cm(0 if lvl == 1 else 0.55)
        pf.space_before = Pt(3 if lvl == 1 else 0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        if page:
            pf.tab_stops.add_tab_stop(Cm(right_cm), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        if i == 0 and open_field:
            r = p.add_run()
            f = OxmlElement("w:fldChar"); f.set(qn("w:fldCharType"), "begin"); f.set(qn("w:dirty"), "true")
            r._r.append(f)
            r2 = p.add_run()
            ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve"); ins.text = FIELD
            r2._r.append(ins)
            r3 = p.add_run()
            sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
            r3._r.append(sep)
        _runs(p, text, size=BODY - 0.5, bold=(lvl == 1), colour=NAVY if lvl == 1 else None)
        if page:                                     # a measured page number, when one is known
            p.add_run("\t" + str(page)).font.size = Pt(BODY - 0.5)
        if i == len(entries) - 1 and close_field:
            rend = OxmlElement("w:fldChar"); rend.set(qn("w:fldCharType"), "end")
            p.add_run()._r.append(rend)


def _update_fields_on_open(doc):
    el = doc.settings.element.find(qn("w:updateFields"))
    if el is None:
        el = OxmlElement("w:updateFields")
        doc.settings.element.append(el)
    el.set(qn("w:val"), "true")


def _footer_pages(doc):
    """Page N in the footer: a field, so whatever the layout engine decides is what gets printed."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for sec in doc.sections:
        f = sec.footer
        f.is_linked_to_previous = False
        p = f.paragraphs[0] if f.paragraphs else f.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in list(p.runs):
            r.text = ""
        a = p.add_run("Page ")
        r = p.add_run()
        b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin"); r._r.append(b)
        c = p.add_run()
        i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = " PAGE "
        c._r.append(i)
        d = p.add_run()
        e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end"); d._r.append(e)
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = GREY


def _figure(doc, rel):
    """one house plate, centred in the column, at exactly the width the page model charged for"""
    path = asset_path(rel)
    if not path.exists():
        raise FileNotFoundError(f"the book asks for a picture that is not there: {rel}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before, pf.space_after = Pt(3), Pt(7)
    p.add_run().add_picture(str(path), width=Cm(IMAGE_WIDTH_CM))   # the width the plan billed

def _page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def _rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:color"), "BBBBBB")
    bdr.append(bot)
    pPr.append(bdr)


def _pack_options(lines, i, width=88):
    """consecutive A) B) C) D) lines are packed into as few lines as a column allows — four stacked
    options turn a two-page practice paper into a six-page one"""
    bits = []
    while i < len(lines) and OPTION.match(lines[i].strip()):
        m = OPTION.match(lines[i].strip())
        bits.append((m.group(1), m.group(2)))
        i += 1
    packed, cur = [], ""
    for L, txt in bits:
        piece = f"{L}) {txt}"
        if cur and len(cur) + 3 + len(piece) > width:
            packed.append(cur)
            cur = piece
        else:
            cur = (cur + "   " + piece).strip()
    if cur:
        packed.append(cur)
    return packed, i


def join_wrapped(lines):
    """markdown folds a sentence across source lines; a Word paragraph must not start where the
    author's editor happened to wrap, so a plain line is glued onto the plain line before it"""
    out = []
    for ln in lines:
        s = ln.strip()
        plain_prev = (out and out[-1].strip() and not re.match(
            r"^\s*(#|\*\*|[•·*-]\s|\(?[A-D]\)|\d+[.)]\s|\|)", out[-1])
            and not out[-1].strip().startswith("---"))
        plain_now = (s and not re.match(r"^\s*(#|\*\*|[•·*-]\s|\(?[A-D]\)|\d+[.)]\s|\||---)", s))
        if out and plain_prev and plain_now:
            out[-1] = out[-1].rstrip() + " " + s
        else:
            out.append(ln)
    return out


def render(md, path, cols=2, gap_cm=0.7):
    """write the .docx, then report the page model of what was just written (see book_layout.plan)"""
    doc = Document()
    _style(doc)
    _landscape_two_col(doc, cols=cols, gap_cm=gap_cm)
    _update_fields_on_open(doc)
    _footer_pages(doc)
    # a two-column A4 page: 29.7cm - 2 x 1.4cm margins - 0.7cm gutter, split between the columns
    col_cm = round((29.7 - 2 * 1.4 - gap_cm) / cols - 0.1, 2)
    flow = flow_of(md)
    pages, brk, problems, info = plan(flow)
    flow, stamped = numbered_flow(flow, pages)          # Contents lines that carry their page number
    title = flow[0][1] if flow and flow[0][0] == "h1" else ""
    i = 0
    while i < len(flow):
        item = flow[i]
        k = item[0]
        if i in brk and doc.paragraphs:
            _page_break(doc)                            # the plan decides where a page ends, exactly
        if k == "h1":
            _heading(doc, item[1], 1, BODY + 6.0, caps=True, before=4, after=4,
                     toc=0 if item[1] == title else 1)
        elif k == "h2" and item[1].strip().lower() == "contents":
            _heading(doc, "Contents", 2, BODY + 3.0, caps=True, before=12)
            entries, j = [], i + 1
            while j < len(flow):
                if flow[j][0] == "toc":
                    entries.append((j, flow[j][1:]))
                elif flow[j][0] == "para" and not entries:
                    _runs(doc.add_paragraph(), flow[j][1], size=BODY - 0.5, italic=True)
                else:
                    break
                j += 1
            # the plan may cut the list across pages: one field, but the page break it asked for
            groups, cur = [], []
            for idx, e in entries:
                if cur and idx in brk:
                    groups.append(cur)
                    cur = []
                cur.append(e)
            if cur:
                groups.append(cur)
            for g, grp in enumerate(groups):
                if g:
                    _page_break(doc)
                _toc(doc, grp, col_cm, open_field=(g == 0), close_field=(g == len(groups) - 1))
            i = j
            continue
        elif k == "h4":
            _heading(doc, item[1], 4, BODY + 0.5, GREY, before=6)
        elif k == "h3":
            _heading(doc, item[1], 3, BODY + 2.0, before=10, toc=2 if is_week(item[1]) else 0)
        elif k == "h2":
            _heading(doc, item[1], 2, BODY + 3.0, caps=True, before=12)
        elif k == "rule":
            _rule(doc)
        elif k == "image":
            _figure(doc, item[1])
        elif k == "opt":
            _indent(doc.add_paragraph(), 1.15, 0.6)
            doc.paragraphs[-1].paragraph_format.space_after = Pt(3)
            _runs(doc.paragraphs[-1], item[1])
        elif k == "num":
            p_ = doc.add_paragraph()
            _indent(p_, 0.62, 0.62)
            _runs(p_, item[1])
        elif k == "bullet":
            p_ = doc.add_paragraph()
            _indent(p_, 0.55, 0.4)
            _runs(p_, item[1])
        elif k == "label":
            p_ = doc.add_paragraph()
            p_.paragraph_format.space_after = Pt(1)
            p_.paragraph_format.keep_with_next = True
            _runs(p_, item[1], bold=True, colour=NAVY)
        else:
            _runs(doc.add_paragraph(), item[1])
        i += 1
    doc.save(path)
    return path, problems, info


if __name__ == "__main__":
    src, dst = sys.argv[1], sys.argv[2]
    _path, problems, info, stamped = render(Path(src).read_text(encoding="utf-8"), Path(dst))
    print(f"wrote {dst} · {info['pages']} pages · {info['breaks']} forced breaks · "
          f"{stamped} Contents lines numbered · problems: {len(problems)}")
    for x in problems[:5]:
        print("   ", x)
